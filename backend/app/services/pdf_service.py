"""
PDF processing service using PyMuPDF and pdfplumber.
Handles all PDF operations including extraction, conversion, and manipulation.
"""

import logging
import fitz  # PyMuPDF
import pdfplumber
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
import io

from app.models.schemas import PDFCommand, OutputFormat
from app.utils.file_manager import FileManager

logger = logging.getLogger(__name__)


class PDFService:
    """
    Core PDF processing service.
    Implements all PDF operations: extraction, conversion, merging, removal.
    """
    
    def __init__(self, file_manager: FileManager):
        """
        Initialize PDF service.
        
        Args:
            file_manager: FileManager instance for file operations
        """
        self.file_manager = file_manager
    
    def get_page_count(self, file_id: str) -> int:
        """
        Get total page count of a PDF.
        
        Args:
            file_id: PDF file identifier
            
        Returns:
            int: Number of pages
        """
        try:
            file_path = self.file_manager.get_file_path(file_id)
            with fitz.open(file_path) as pdf:
                return pdf.page_count
        except Exception as e:
            logger.error(f"Failed to get page count: {e}")
            raise ValueError(f"Cannot read PDF: {e}")
    
    def process_command(self, file_id: str, command: PDFCommand) -> str:
        """
        Execute PDF processing command.
        
        Args:
            file_id: Source PDF file identifier
            command: Parsed command to execute
            
        Returns:
            str: Result file_id
        """
        try:
            action_map = {
                "convert_to_word": self._convert_to_word,
                "extract_pages": self._extract_pages,
                "extract_page_range": self._extract_page_range,
                "remove_pages": self._remove_pages,
                "merge_pages": self._merge_pages,
                "extract_to_word": self._extract_to_word,
            }
            
            if command.action not in action_map:
                raise ValueError(f"Unsupported action: {command.action}")
            
            action_func = action_map[command.action]
            result_file_id = action_func(file_id, command)
            
            logger.info(f"Command executed: {command.action} -> {result_file_id}")
            return result_file_id
            
        except Exception as e:
            logger.error(f"Command execution failed: {e}")
            raise
    
    def _convert_to_word(self, file_id: str, command: PDFCommand) -> str:
        """
        Convert entire PDF to Word document.
        Extracts text and preserves basic formatting.
        """
        try:
            file_path = self.file_manager.get_file_path(file_id)
            file_info = self.file_manager.get_file_info(file_id)
            
            # Create Word document
            doc = Document()
            
            # Process PDF with pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Add page heading
                    doc.add_heading(f'Page {page_num}', level=2)
                    
                    # Extract text
                    text = page.extract_text()
                    if text:
                        # Add text content
                        for paragraph in text.split('\n\n'):
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            if table_data:
                                self._add_table_to_doc(doc, table_data)
                    
                    # Page break (except last page)
                    if page_num < len(pdf.pages):
                        doc.add_page_break()
            
            # Save Word document
            output_path = self.file_manager.get_temp_path(suffix=".docx")
            doc.save(output_path)
            
            # Register result file
            base_name = Path(file_info['filename']).stem
            result_filename = f"{base_name}_converted.docx"
            result_file_id = self.file_manager.save_result(
                output_path,
                result_filename,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Cleanup temp file
            output_path.unlink()
            
            return result_file_id
            
        except Exception as e:
            logger.error(f"PDF to Word conversion failed: {e}")
            raise ValueError(f"Conversion failed: {e}")
    
    def _extract_pages(self, file_id: str, command: PDFCommand) -> str:
        """Extract specific pages from PDF."""
        try:
            file_path = self.file_manager.get_file_path(file_id)
            file_info = self.file_manager.get_file_info(file_id)
            
            pages_to_extract = command.pages
            if not isinstance(pages_to_extract, list):
                pages_to_extract = [pages_to_extract]
            
            # Validate page numbers
            total_pages = self.get_page_count(file_id)
            if any(p > total_pages for p in pages_to_extract):
                raise ValueError(f"Page number exceeds document length ({total_pages} pages)")
            
            # Open source PDF
            source_pdf = fitz.open(file_path)
            
            # Create new PDF with extracted pages
            output_pdf = fitz.open()
            for page_num in pages_to_extract:
                # PyMuPDF uses 0-based indexing
                output_pdf.insert_pdf(source_pdf, from_page=page_num - 1, to_page=page_num - 1)
            
            # Save extracted PDF
            output_path = self.file_manager.get_temp_path(suffix=".pdf")
            output_pdf.save(output_path)
            output_pdf.close()
            source_pdf.close()
            
            # Convert to Word if requested
            if command.output_format == OutputFormat.DOCX:
                result_file_id = self._convert_pdf_to_word(
                    output_path,
                    f"{Path(file_info['filename']).stem}_pages_{'_'.join(map(str, pages_to_extract))}.docx"
                )
                output_path.unlink()
            else:
                base_name = Path(file_info['filename']).stem
                pages_str = '_'.join(map(str, pages_to_extract))
                result_filename = f"{base_name}_pages_{pages_str}.pdf"
                result_file_id = self.file_manager.save_result(
                    output_path,
                    result_filename,
                    mime_type="application/pdf"
                )
                output_path.unlink()
            
            return result_file_id
            
        except Exception as e:
            logger.error(f"Page extraction failed: {e}")
            raise ValueError(f"Extraction failed: {e}")
    
    def _extract_page_range(self, file_id: str, command: PDFCommand) -> str:
        """Extract a range of pages from PDF."""
        try:
            file_path = self.file_manager.get_file_path(file_id)
            file_info = self.file_manager.get_file_info(file_id)
            
            start_page = command.page_range["start"]
            end_page = command.page_range["end"]
            
            # Validate page range
            total_pages = self.get_page_count(file_id)
            if end_page > total_pages:
                raise ValueError(f"End page {end_page} exceeds document length ({total_pages} pages)")
            
            # Open source PDF
            source_pdf = fitz.open(file_path)
            
            # Create new PDF with page range (0-based indexing)
            output_pdf = fitz.open()
            output_pdf.insert_pdf(source_pdf, from_page=start_page - 1, to_page=end_page - 1)
            
            # Save extracted PDF
            output_path = self.file_manager.get_temp_path(suffix=".pdf")
            output_pdf.save(output_path)
            output_pdf.close()
            source_pdf.close()
            
            # Convert to Word if requested
            if command.output_format == OutputFormat.DOCX:
                result_file_id = self._convert_pdf_to_word(
                    output_path,
                    f"{Path(file_info['filename']).stem}_pages_{start_page}-{end_page}.docx"
                )
                output_path.unlink()
            else:
                base_name = Path(file_info['filename']).stem
                result_filename = f"{base_name}_pages_{start_page}-{end_page}.pdf"
                result_file_id = self.file_manager.save_result(
                    output_path,
                    result_filename,
                    mime_type="application/pdf"
                )
                output_path.unlink()
            
            return result_file_id
            
        except Exception as e:
            logger.error(f"Page range extraction failed: {e}")
            raise ValueError(f"Extraction failed: {e}")
    
    def _remove_pages(self, file_id: str, command: PDFCommand) -> str:
        """Remove specified pages from PDF."""
        try:
            file_path = self.file_manager.get_file_path(file_id)
            file_info = self.file_manager.get_file_info(file_id)
            
            pages_to_remove = command.pages
            if not isinstance(pages_to_remove, list):
                pages_to_remove = [pages_to_remove]
            
            # Validate page numbers
            total_pages = self.get_page_count(file_id)
            if any(p > total_pages for p in pages_to_remove):
                raise ValueError(f"Page number exceeds document length ({total_pages} pages)")
            
            # Convert to 0-based and sort in reverse to avoid index shifting
            pages_to_remove_idx = sorted([p - 1 for p in pages_to_remove], reverse=True)
            
            # Open PDF
            pdf_doc = fitz.open(file_path)
            
            # Delete pages
            for page_idx in pages_to_remove_idx:
                pdf_doc.delete_page(page_idx)
            
            # Save modified PDF
            output_path = self.file_manager.get_temp_path(suffix=".pdf")
            pdf_doc.save(output_path)
            pdf_doc.close()
            
            # Register result
            base_name = Path(file_info['filename']).stem
            pages_str = '_'.join(map(str, pages_to_remove))
            result_filename = f"{base_name}_removed_pages_{pages_str}.pdf"
            result_file_id = self.file_manager.save_result(
                output_path,
                result_filename,
                mime_type="application/pdf"
            )
            
            # Cleanup
            output_path.unlink()
            
            return result_file_id
            
        except Exception as e:
            logger.error(f"Page removal failed: {e}")
            raise ValueError(f"Removal failed: {e}")
    
    def _merge_pages(self, file_id: str, command: PDFCommand) -> str:
        """
        Merge specified pages into a new PDF.
        Similar to extract_pages but emphasizes merging.
        """
        return self._extract_pages(file_id, command)
    
    def _extract_to_word(self, file_id: str, command: PDFCommand) -> str:
        """Extract pages and convert to Word."""
        # Set output format to DOCX
        command.output_format = OutputFormat.DOCX
        
        # Use extract_pages which handles DOCX conversion
        return self._extract_pages(file_id, command)
    
    def _convert_pdf_to_word(self, pdf_path: Path, output_filename: str) -> str:
        """
        Helper to convert a PDF file to Word.
        
        Args:
            pdf_path: Path to PDF file
            output_filename: Desired output filename
            
        Returns:
            str: Result file_id
        """
        try:
            doc = Document()
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    doc.add_heading(f'Page {page_num}', level=2)
                    
                    text = page.extract_text()
                    if text:
                        for paragraph in text.split('\n\n'):
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())
                    
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            if table_data:
                                self._add_table_to_doc(doc, table_data)
                    
                    if page_num < len(pdf.pages):
                        doc.add_page_break()
            
            # Save Word document
            word_path = self.file_manager.get_temp_path(suffix=".docx")
            doc.save(word_path)
            
            # Register result
            result_file_id = self.file_manager.save_result(
                word_path,
                output_filename,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            word_path.unlink()
            return result_file_id
            
        except Exception as e:
            logger.error(f"PDF to Word conversion failed: {e}")
            raise ValueError(f"Conversion failed: {e}")
    
    def _add_table_to_doc(self, doc: Document, table_data: List[List]) -> None:
        """
        Add extracted table to Word document.
        
        Args:
            doc: Document object
            table_data: 2D list of table cells
        """
        try:
            if not table_data or not table_data[0]:
                return
            
            # Create table
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            # Fill table
            for i, row_data in enumerate(table_data):
                for j, cell_value in enumerate(row_data):
                    if j < cols and cell_value:
                        table.rows[i].cells[j].text = str(cell_value)
            
            doc.add_paragraph()  # Add spacing after table
            
        except Exception as e:
            logger.warning(f"Failed to add table to document: {e}")
